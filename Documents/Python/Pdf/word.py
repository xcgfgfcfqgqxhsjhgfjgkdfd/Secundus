import docx

#reading
W=input("Entrez le chemin d'acces : ")
doc = docx.Document(W)
all_paras = doc.paragraphs
print(len(all_paras))

for para in all_paras:
    print(para.text)
    print("-------")

# code for a single paragraph-print
single_para = doc.paragraphs[4]
for run in single_para.runs:
    print(run.text)

