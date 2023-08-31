import docx

#Writing
mydoc = docx.Document()
mydoc.add_paragraph("This is first paragraph of a MS Word file.")
third_para = mydoc.add_paragraph("This is the third paragraph.")
third_para.add_run(" this is a section at the end of third paragraph")
mydoc.add_heading("This is level 1 heading", 0)
mydoc.add_heading("This is level 2 heading", 1)
mydoc.add_heading("This is level 3 heading", 2)
mydoc.add_picture("Word.png", width=docx.shared.Inches(5), height=docx.shared.Inches(7))
S = input("Entrez le chemin de sauvegarde du fichier : ")
mydoc.save(S)